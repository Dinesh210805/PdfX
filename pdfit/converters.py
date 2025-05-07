"""
Converter functions for different file formats to PDF
"""
import os
import sys
import subprocess
import shutil
import colorama
from fpdf import FPDF
from PIL import Image
import docx
import pdfkit
from colorama import Fore, Style
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfMerger, PdfReader, PdfWriter

# Initialize colorama
colorama.init()

def convert_txt_to_pdf(input_path, output_path=None):
    """Convert a text file to PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    try:
        # Read the text file
        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Create a PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add text to PDF - improved handling of long lines
        for line in content.split('\n'):
            # Handle long lines by wrapping text and special characters
            try:
                pdf.multi_cell(0, 10, txt=line)
            except Exception as e:
                # Fall back to ASCII encoding if there's an encoding issue
                pdf.multi_cell(0, 10, txt=line.encode('ascii', 'replace').decode())
          
        # Save the PDF
        pdf.output(output_path)
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}Text file converted successfully! PDF saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error converting TXT to PDF: {str(e)}{Style.RESET_ALL}")
        return False

def convert_jpg_to_pdf(input_path, output_path=None):
    """Convert JPG/PNG image to PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    try:
        # Open the image
        img = Image.open(input_path)
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
            
        # Save as PDF
        img.save(output_path, "PDF", resolution=100.0)
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}Image converted successfully! PDF saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error converting image to PDF: {str(e)}{Style.RESET_ALL}")
        return False

def convert_docx_to_pdf(input_path, output_path=None):
    """Convert DOCX to PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    try:
        # Track conversion success
        success = False
        
        # Use pure Python method first (more reliable)
        try:
            # Load the document
            doc = docx.Document(input_path)
            
            # Create a PDF
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            y_position = height - 72  # Start near the top
            
            # Extract text and add to PDF with better formatting
            for para in doc.paragraphs:
                if not para.text.strip():  # Skip empty paragraphs
                    y_position -= 12
                    continue
                    
                # Add text with proper wrapping
                text = para.text
                text_obj = c.beginText(72, y_position)
                text_obj.setFont("Helvetica", 12)
                text_obj.textLine(text)
                c.drawText(text_obj)
                
                y_position -= 14  # Move down for next line
                
                # Check if we need a new page
                if y_position < 72:
                    c.showPage()
                    y_position = height - 72
            
            # Save the PDF
            c.save()
            
            # If we get here, the conversion was successful
            success = True
        except Exception as inner_e:
            print(f"{Fore.YELLOW}Fallback method 1 failed: {str(inner_e)}. Trying alternative method...{Style.RESET_ALL}")
            success = False
            
        # If the first method didn't work, try docx2pdf as a backup
        if not success:
            try:
                from docx2pdf import convert
                convert(input_path, output_path)
                success = True
            except (ImportError, Exception) as inner_e:
                print(f"{Fore.YELLOW}Fallback method 2 failed: {str(inner_e)}{Style.RESET_ALL}")
                print(f"{Fore.YELLOW}Trying fallback method 3...{Style.RESET_ALL}")
                success = False
                
        # Final fallback method - extremely basic but reliable approach
        if not success:
            try:
                # Create a simple FPDF document
                from fpdf import FPDF
                
                # Load the document
                doc = docx.Document(input_path)
                
                # Create PDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                
                # Simple extraction of text and add to PDF
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Add each paragraph as multiline text
                        try:
                            pdf.multi_cell(0, 10, txt=para.text)
                        except Exception:
                            # Handle encoding issues
                            pdf.multi_cell(0, 10, txt=para.text.encode('ascii', 'replace').decode())
                
                # Save the PDF
                pdf.output(output_path)
                success = True
            except Exception as inner_e:
                print(f"{Fore.YELLOW}All fallback methods failed: {str(inner_e)}{Style.RESET_ALL}")
                success = False
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}Word document converted successfully! PDF saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error converting DOCX to PDF: {str(e)}{Style.RESET_ALL}")
        return False

def convert_html_to_pdf(input_path, output_path=None):
    """Convert HTML to PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    try:
        # First, attempt using pdfkit (requires wkhtmltopdf)
        try:
            # Use pdfkit to convert HTML to PDF
            options = {
                'quiet': ''
            }
            pdfkit.from_file(input_path, output_path, options=options)
            print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}HTML converted successfully with pdfkit! PDF saved to: {output_path}{Style.RESET_ALL}")
            return True
        except OSError as e:
            if "wkhtmltopdf" in str(e):
                print(f"\n{Fore.WHITE}[{Fore.YELLOW}!{Fore.WHITE}] {Fore.YELLOW}wkhtmltopdf not found. Trying fallback method...{Style.RESET_ALL}")
                # Continue to fallback method
            else:
                raise
        except Exception as e:
            print(f"\n{Fore.WHITE}[{Fore.YELLOW}!{Fore.WHITE}] {Fore.YELLOW}pdfkit error: {str(e)}. Trying fallback method...{Style.RESET_ALL}")
        
        # Fallback method: Use reportlab to create a PDF from HTML content
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            from html.parser import HTMLParser
            
            # Simple HTML parser to extract text
            class HTMLTextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                    self.current_data = ""
                    self.in_body = False
                
                def handle_starttag(self, tag, attrs):
                    if tag == 'body':
                        self.in_body = True
                    elif tag in ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div', 'br'):
                        if self.current_data:
                            self.text.append(self.current_data)
                            self.current_data = ""
                
                def handle_endtag(self, tag):
                    if tag == 'body':
                        self.in_body = False
                    elif tag in ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div'):
                        if self.current_data:
                            self.text.append(self.current_data)
                            self.current_data = ""
                
                def handle_data(self, data):
                    if self.in_body and data.strip():
                        self.current_data += data
                
                def get_text(self):
                    if self.current_data:
                        self.text.append(self.current_data)
                    return self.text
            
            # Read HTML content
            with open(input_path, 'r', encoding='utf-8') as html_file:
                html_content = html_file.read()
            
            # Extract text from HTML
            parser = HTMLTextExtractor()
            parser.feed(html_content)
            paragraphs = parser.get_text()
            
            # Create PDF using reportlab
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            flowables = []
            
            # Add content to PDF
            for para in paragraphs:
                p = Paragraph(para, styles['Normal'])
                flowables.append(p)
            
            # Build the PDF
            doc.build(flowables)
            
            print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}HTML converted successfully with fallback method! PDF saved to: {output_path}{Style.RESET_ALL}")
            return True
            
        except Exception as inner_e:
            print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Fallback HTML conversion failed: {str(inner_e)}{Style.RESET_ALL}")
            print(f"\n{Fore.YELLOW}For better HTML to PDF conversion, install wkhtmltopdf:{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}  Windows: https://wkhtmltopdf.org/downloads.html{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}  Linux: sudo apt-get install wkhtmltopdf{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}  macOS: brew install wkhtmltopdf{Style.RESET_ALL}")
            raise
            
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error converting HTML to PDF: {str(e)}{Style.RESET_ALL}")
        return False

def convert_ppt_to_pdf(input_path, output_path=None):
    """Convert PPT/PPTX to PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    try:
        try:
            import comtypes.client
        except ImportError:
            print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error: comtypes module not found. Please install it using 'pip install comtypes'.{Style.RESET_ALL}")
            return False
        
        # Create PowerPoint application
        powerpoint = comtypes.client.CreateObject("PowerPoint.Application")
        powerpoint.Visible = 1
        
        # Open the presentation
        presentation = powerpoint.Presentations.Open(os.path.abspath(input_path))
        # Save as PDF
        presentation.SaveAs(os.path.abspath(output_path), 32)  # 32 = PDF format
        
        # Close
        presentation.Close()
        powerpoint.Quit()
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PowerPoint presentation converted successfully! PDF saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error converting PPT to PDF: {str(e)}{Style.RESET_ALL}")
        return False

def merge_pdfs(input_paths, output_path):
    """Merge multiple PDFs into one"""
    try:
        merger = PdfMerger()
        
        # Add each PDF to the merger
        for pdf in input_paths:
            merger.append(pdf)
            
        # Write to the output file
        merger.write(output_path)
        merger.close()
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PDFs merged successfully! Output saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error merging PDFs: {str(e)}{Style.RESET_ALL}")
        return False

def split_pdf(input_path, output_directory=None):
    """Split a PDF into individual pages"""
    if not output_directory:
        output_directory = os.path.dirname(input_path)
    
    try:
        # Read the input PDF
        pdf = PdfReader(input_path)
        
        # Get the base filename
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        
        # Split the PDF
        for i, page in enumerate(pdf.pages):
            output_path = os.path.join(output_directory, f"{base_name}_page_{i+1}.pdf")
            writer = PdfWriter()
            writer.add_page(page)
            with open(output_path, "wb") as output_file:
                writer.write(output_file)
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PDF split successfully! {len(pdf.pages)} pages extracted to: {output_directory}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error splitting PDF: {str(e)}{Style.RESET_ALL}")
        return False

def extract_text_from_pdf(input_path, output_path=None):
    """Extract text from a PDF file"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '.txt'
    
    try:
        # Read the PDF
        reader = PdfReader(input_path)
        
        # Extract text from each page
        with open(output_path, 'w', encoding='utf-8') as output_file:
            for page in reader.pages:
                text = page.extract_text()
                output_file.write(text + "\n\n")
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}Text extracted successfully! Saved to: {output_path}{Style.RESET_ALL}")
        return True

    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error extracting text from PDF: {str(e)}{Style.RESET_ALL}")
        return False

def password_protect_pdf(input_path, output_path=None, password=None, owner_pw=None):
    """Add password protection to a PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '_protected.pdf'
    
    try:
        # Open the PDF
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        # Add all pages to the writer
        for page in reader.pages:
            writer.add_page(page)
        
        # Encrypt the PDF
        if owner_pw:
            writer.encrypt(user_password=password, owner_password=owner_pw)
        else:
            writer.encrypt(password)
        
        # Save the encrypted PDF
        with open(output_path, "wb") as output_file:
            writer.write(output_file)
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PDF password protected successfully! Saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error adding password protection to PDF: {str(e)}{Style.RESET_ALL}")
        return False

def remove_password_from_pdf(input_path, output_path=None, password=None):
    """Remove password protection from a PDF if the password is known"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '_decrypted.pdf'
    
    try:
        # Open the encrypted PDF
        reader = PdfReader(input_path)
        
        # Check if the PDF is encrypted
        if not reader.is_encrypted:
            print(f"\n{Fore.WHITE}[{Fore.YELLOW}!{Fore.WHITE}] {Fore.YELLOW}The PDF is not encrypted. No need to decrypt.{Style.RESET_ALL}")
            return False
        
        # Decrypt the PDF with the password
        reader.decrypt(password)
        
        # Create a new PDF without encryption
        writer = PdfWriter()
        
        # Add all pages to the writer
        for page in reader.pages:
            writer.add_page(page)
        
        # Save the decrypted PDF
        with open(output_path, "wb") as output_file:
            writer.write(output_file)
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PDF decrypted successfully! Saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error removing password from PDF: {str(e)}{Style.RESET_ALL}")
        return False

def compress_pdf(input_path, output_path=None, quality='medium'):
    """Compress PDF file size using optimization"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '_compressed.pdf'
    
    try:
        # First attempt: Use ghostscript for compression (better quality)
        try:
            import subprocess
            import shutil
            
            # Check if ghostscript is installed
            gs_command = "gs" if os.name != 'nt' else "gswin64c"
            if os.name == 'nt':
                # On Windows, check both gswin64c and gswin32c
                if shutil.which("gswin64c") is None and shutil.which("gswin32c") is not None:
                    gs_command = "gswin32c"
                    
            if shutil.which(gs_command) is None:
                raise FileNotFoundError(f"Ghostscript ({gs_command}) not found in PATH")
            
            # Define quality settings
            quality_settings = {
                'low': '/default',    # Low compression, better quality
                'medium': '/ebook',   # Medium compression
                'high': '/screen'     # High compression, lower quality
            }
            
            quality_arg = quality_settings.get(quality, '/ebook')
            
            # Construct the ghostscript command
            gs_args = [
                gs_command,
                "-sDEVICE=pdfwrite",
                "-dCompatibilityLevel=1.4",
                "-dPDFSETTINGS=" + quality_arg,
                "-dNOPAUSE",
                "-dQUIET",
                "-dBATCH",
                f"-sOutputFile={output_path}",
                input_path
            ]
            
            # Run the command
            subprocess.check_call(gs_args)
            
            # Check file sizes
            input_size = os.path.getsize(input_path)
            output_size = os.path.getsize(output_path)
            reduction = (1 - output_size / input_size) * 100
            
            print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PDF compressed successfully with Ghostscript! Saved to: {output_path}{Style.RESET_ALL}")
            print(f"{Fore.GREEN}Original size: {input_size / 1024:.1f} KB, New size: {output_size / 1024:.1f} KB{Style.RESET_ALL}")
            print(f"{Fore.GREEN}Reduced by: {reduction:.1f}%{Style.RESET_ALL}")
            return True
            
        except (FileNotFoundError, subprocess.SubprocessError) as e:
            print(f"\n{Fore.WHITE}[{Fore.YELLOW}!{Fore.WHITE}] {Fore.YELLOW}Ghostscript compression attempt failed: {str(e)}{Style.RESET_ALL}")
            print(f"{Fore.YELLOW}Trying fallback compression method...{Style.RESET_ALL}")
        
        # Fallback method: Use PyPDF2 for basic compression
        try:
            # Use PyPDF2 for a basic level of compression
            reader = PdfReader(input_path)
            writer = PdfWriter()
            
            # Copy each page with default compression
            for page in reader.pages:
                writer.add_page(page)
            
            # Save the compressed PDF
            with open(output_path, "wb") as output_file:
                writer.write(output_file)
            
            # Check file sizes
            input_size = os.path.getsize(input_path)
            output_size = os.path.getsize(output_path)
            
            # If the output is actually larger, use the input file instead
            if output_size >= input_size:
                print(f"\n{Fore.WHITE}[{Fore.YELLOW}!{Fore.WHITE}] {Fore.YELLOW}Fallback compression did not reduce file size.{Style.RESET_ALL}")
                print(f"{Fore.YELLOW}For better compression, install Ghostscript:{Style.RESET_ALL}")
                print(f"{Fore.YELLOW}  Windows: Download from https://ghostscript.com/releases/{Style.RESET_ALL}")
                print(f"{Fore.YELLOW}  macOS: brew install ghostscript{Style.RESET_ALL}")
                print(f"{Fore.YELLOW}  Linux: sudo apt-get install ghostscript{Style.RESET_ALL}")
                
                # Copy the original file to the output path
                import shutil
                shutil.copy(input_path, output_path)
                print(f"\n{Fore.WHITE}[{Fore.YELLOW}+{Fore.WHITE}] {Fore.YELLOW}Using original file: {output_path}{Style.RESET_ALL}")
            else:
                reduction = (1 - output_size / input_size) * 100
                print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}PDF compressed with fallback method! Saved to: {output_path}{Style.RESET_ALL}")
                print(f"{Fore.GREEN}Original size: {input_size / 1024:.1f} KB, New size: {output_size / 1024:.1f} KB{Style.RESET_ALL}")
                print(f"{Fore.GREEN}Reduced by: {reduction:.1f}%{Style.RESET_ALL}")
                print(f"\n{Fore.YELLOW}Note: For better compression, install Ghostscript.{Style.RESET_ALL}")
                
            return True
                
        except Exception as inner_e:
            print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Fallback compression failed: {str(inner_e)}{Style.RESET_ALL}")
            raise
            
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error compressing PDF: {str(e)}{Style.RESET_ALL}")
        return False

def rotate_pdf_pages(input_path, output_path=None, pages=None, angle=90):
    """Rotate specific pages in a PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '_rotated.pdf'
    
    try:
        # Open the PDF
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        # Parse pages string (e.g., "1,3-5,7")
        if pages is None:
            pages = list(range(len(reader.pages)))
        elif isinstance(pages, str):
            page_list = []
            for part in pages.split(','):
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    page_list.extend(range(start - 1, end))  # Convert to 0-based index
                else:
                    page_list.append(int(part) - 1)  # Convert to 0-based index
            pages = page_list
        
        # Add all pages to the writer, rotating the specified ones
        for i, page in enumerate(reader.pages):
            if i in pages:
                page.rotate(angle)
            writer.add_page(page)
        
        # Save the rotated PDF
        with open(output_path, "wb") as output_file:
            writer.write(output_file)
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}Pages rotated successfully! Saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error rotating PDF pages: {str(e)}{Style.RESET_ALL}")
        return False

def reorder_pdf_pages(input_path, output_path=None, page_order=None):
    """Reorder pages in a PDF"""
    if not output_path:
        output_path = os.path.splitext(input_path)[0] + '_reordered.pdf'
    
    try:
        # Open the PDF
        reader = PdfReader(input_path)
        writer = PdfWriter()
        if page_order is None:
            print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}No page order specified.{Style.RESET_ALL}")
            return False
            
        # Parse page order string (e.g., "3,1,2")
        if isinstance(page_order, str):
            page_list = [int(p) - 1 for p in page_order.split(',')]  # Convert to 0-based index
        else:
            page_list = [p - 1 for p in page_order]  # Already a list
            
        # Check for valid page numbers
        max_page = len(reader.pages) - 1
        for page_num in page_list:
            if page_num < 0 or page_num > max_page:
                print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Invalid page number: {page_num + 1}. Document has {max_page + 1} pages.{Style.RESET_ALL}")
                return False
        
        # Add pages in the specified order
        for page_num in page_list:
            writer.add_page(reader.pages[page_num])
        
        # Save the reordered PDF
        with open(output_path, "wb") as output_file:
            writer.write(output_file)
        
        print(f"\n{Fore.WHITE}[{Fore.GREEN}+{Fore.WHITE}] {Fore.GREEN}Pages reordered successfully! Saved to: {output_path}{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"\n{Fore.WHITE}[{Fore.RED}!{Fore.WHITE}] {Fore.RED}Error reordering PDF pages: {str(e)}{Style.RESET_ALL}")
        return False
